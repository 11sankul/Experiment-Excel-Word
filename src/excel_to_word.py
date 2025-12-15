import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, ns

# ---------------- CONFIG ----------------
EXCEL_PATH = "data/input.xlsx"
SHEET_NAME = "Sheet1"
OUTPUT_PATH = "output/output.docx"

# ---------------- LOAD EXCEL ----------------
df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME)
df = df.fillna("")  # blanks stay blanks

# ---------------- CREATE DOCUMENT ----------------
document = Document()
section = document.sections[0]

# ---------------- SET EQUAL MARGINS ----------------
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

# ---------------- HEADER (SAME FOR ALL PAGES) ----------------
header = section.header
header.paragraphs[0].clear()
hp = header.paragraphs[0]

# Left
hp.add_run("School report").font.size = Pt(10)
hp.add_run("\t")

# Center
hp.add_run("Sankul").font.size = Pt(10)
hp.add_run("\t")

# Right
hp.add_run("Moshi-Chikhali-Talwade-Dehu-Dehu Road").font.size = Pt(10)

# ---------------- FOOTER (PAGE NUMBERS ONLY) ----------------
footer = section.footer
footer.paragraphs[0].clear()
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = fp.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(ns.qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(ns.qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# ---------------- PAGE BORDER ----------------
def add_page_border(sec):
    sectPr = sec._sectPr
    pgBorders = OxmlElement('w:pgBorders')

    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '8')      # ~1pt
        border.set(ns.qn('w:space'), '24')
        border.set(ns.qn('w:color'), '000000')
        pgBorders.append(border)

    sectPr.append(pgBorders)

add_page_border(section)

# ---------------- CONTENT PER EXCEL ROW ----------------
for idx, row in df.iterrows():

    # EXACT school name from THIS Excel row
    school_name = row.get("Name of School", "")

    # SINGLE BODY HEADING (NO DUPLICATE)
    p = document.add_paragraph(f"School Name – {school_name}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    # TABLE WITH ALL COLUMNS
    table = document.add_table(rows=len(df.columns), cols=2)
    table.style = "Table Grid"

    for i, col in enumerate(df.columns):
        table.cell(i, 0).text = str(col)
        table.cell(i, 1).text = str(row[col])

    if idx < len(df) - 1:
        document.add_page_break()

# ---------------- SAVE ----------------
document.save(OUTPUT_PATH)
print("Word file generated exactly as specified.")
